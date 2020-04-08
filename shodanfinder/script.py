from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE
from docx import Document
from shodan_finder import ShodanFinder
from datetime import datetime

shodan_instance = ShodanFinder(['adacis.net', 'google.com', 'facebook.net'])

sites_services = shodan_instance.shodanProcedure()

doc = Document()
doc.add_heading('Module X : Shodan Finder', 0)

# Custom Style
obj_styles = doc.styles
obj_charstyle = obj_styles.add_style('Banner', WD_STYLE_TYPE.CHARACTER)
obj_font = obj_charstyle.font
obj_font.size = Pt(10)
obj_font.name = 'Montserrat'

for site_service in sites_services:
    doc.add_paragraph("•  Site Name : {}".format(site_service['name']))
    doc.add_paragraph("     >  Host : {}".format(site_service['host']))

    for service in site_service['services']:
        doc.add_paragraph("     >  Service at port : {} ".format(service['port']))
        doc.add_paragraph("         +  serveur : {} ".format(service['serveur']))
        doc.add_paragraph("         +  technologies : {} ".format(service['technologies']))
        
        doc.add_paragraph("         +  banner : ")
        doc.add_paragraph("")
        
        paragraph = doc.add_paragraph("")
        paragraph_format = paragraph.paragraph_format
        paragraph_format.line_spacing = Pt(18)
        paragraph.add_run(service['banner'], style = 'Banner')

date = datetime.now()
doc.save("results/shodan_finder_result_{}.docx".format(date.strftime("%m%d%Y%H%M%S")))

print('The document has been generated successfully')