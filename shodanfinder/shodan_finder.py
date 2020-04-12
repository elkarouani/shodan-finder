from re import match
import shodan
from socket import gethostbyname, gaierror
from docx.shared import Inches, Pt
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL
from docx import Document
from datetime import datetime

try:
    from secret import SHODAN_API_KEY
except :
    pass

class ShodanFinder :
    def __init__(self, sites):
        self.sites_list = sites
        self.data_instance = {}
        self.shodan_key = SHODAN_API_KEY if 'SHODAN_API_KEY' in globals() else ""

    def getHost(self, site):
        if match(r'[a-zA-Z0-9]*.[a-zA-Z0-9]*', site):
            return gethostbyname(site)
        else:
            raise gaierror

    def setAPIKey(self, key):
        self.shodan_key = key

    def extractData(self, host):
        api = shodan.Shodan(self.shodan_key)
        self.data_instance = api.host(host)

    def extractPortsList(self):
        return self.data_instance['ports']
    
    def extractUsedServer(self, port):
        for port_container in self.data_instance['data']:
            if port_container['port'] == port:
                if not 'http' in port_container or not 'server' in port_container['http'] : return "Not Found"

                return port_container['http']['server']

    def extractBannere(self, port):
        for port_container in self.data_instance['data']:
            if port_container['port'] == port:
                return port_container['data']

    def extractTechnologies(self, port):
        for port_container in self.data_instance['data']:
            if port_container['port'] == port:
                if not 'http' in port_container or not 'components' in port_container['http'] : return "Not Found"
                return port_container['http']['components'] if port_container['http']['components'] != {} else "Not Found"

    def shodanProcedure(self):
        sites_services = []

        for site in self.sites_list:
            try:
                host = self.getHost(site)
                sites_services.append({'name': site})

                self.extractData(host)
                sites_services[len(sites_services) - 1]['host'] = host

                ports = self.extractPortsList()
                services = []
                
                if ports == [] : services.append({'ports': "Not Found"})

                for port in ports:
                    services.append({
                        'port': port,
                        'serveur': self.extractUsedServer(port),
                        'banner': self.extractBannere(port),
                        'technologies': self.extractTechnologies(port)
                    })
                sites_services[len(sites_services) - 1]['services'] = services
            except gaierror:
                print(f"\n • The address requested ['{site}'] is not associated with any host. \n")
                pass
        

        return sites_services

    def printDocument(self, sites_services):
        document_instance = Document()
        document_instance.add_heading('Module X : Shodan Finder', 0)

        table = document_instance.add_table(rows=1, cols=5)
        table.autofit = False
        table.style = 'Table Grid'

        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Website'
        hdr_cells[1].text = 'Host'
        hdr_cells[2].text = 'Ports'
        hdr_cells[3].text = 'Serveur'
        hdr_cells[4].text = 'Technologies'

        for site_service in sites_services: 
            site_service_website_cells_to_merge = []
            site_service_host_cells_to_merge = []

            
            for service in site_service['services']: 
                new_row_1 = table.add_row()
                new_row_2 = table.add_row()

                new_row_2.cells[2].merge(new_row_2.cells[3])
                new_row_2.cells[3].merge(new_row_2.cells[4])

                site_service_website_cells_to_merge.append(new_row_1.cells[0])
                site_service_website_cells_to_merge.append(new_row_2.cells[0])
                site_service_host_cells_to_merge.append(new_row_1.cells[1])
                site_service_host_cells_to_merge.append(new_row_2.cells[1])
            
            for index in range(len(site_service_website_cells_to_merge) - 1):
                current_cell = site_service_website_cells_to_merge[index]
                next_cell = site_service_website_cells_to_merge[index + 1]
                current_cell.merge(next_cell)
            
            for index in range(len(site_service_host_cells_to_merge) - 1):
                current_cell = site_service_host_cells_to_merge[index]
                next_cell = site_service_host_cells_to_merge[index + 1]
                current_cell.merge(next_cell)        

        row_start = 1
        for index in range(len(sites_services)):
            row_instance = table.rows[row_start].cells
            row_instance[0].text = sites_services[index]['name']
            row_instance[0].vertical_alignment = WD_ALIGN_VERTICAL.TOP
            row_instance[1].text = sites_services[index]['host']
            row_instance[1].vertical_alignment = WD_ALIGN_VERTICAL.TOP

            for service in sites_services[index]['services']:
                row_instance[2].text = str(service['port'])
                row_instance[2].vertical_alignment = WD_ALIGN_VERTICAL.TOP
                row_instance[3].text = str(service['serveur'])
                row_instance[3].vertical_alignment = WD_ALIGN_VERTICAL.TOP

                if service['technologies'] != "Not Found" :
                    for technologie in service['technologies']:
                        technologie_type = service['technologies'][technologie]['categories'][0]

                        row_instance[4].add_paragraph(f'• {technologie_type} :')
                        row_instance[4].add_paragraph(f'- {technologie}')
                else :
                    row_instance[4].text = str(service['technologies'])
                    row_instance[4].vertical_alignment = WD_ALIGN_VERTICAL.TOP

                row_start += 1
                row_instance = table.rows[row_start].cells if row_start < len(table.rows) else None

                row_instance[2].text = str(service['banner'])
                row_instance[2].vertical_alignment = WD_ALIGN_VERTICAL.TOP

                row_start += 1
                row_instance = table.rows[row_start].cells if row_start < len(table.rows) else None


        date = datetime.now()
        document_instance.save("results/shodan_finder_result_{}.docx".format(date.strftime("%m%d%Y%H%M%S")))

        print('The document has been generated successfully')